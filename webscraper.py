import urllib.request
import time
from docx import Document
from bs4 import BeautifulSoup
import requests
from pprint import pprint
import nltk
from typing import List, Any, Tuple
from collections import Counter


# Uncomment the lines below if you get an error
# on the first run of the program.
# After that you can delete them or comment them again.
# nltk.download('punkt')
# nltk.download('stopwords')


class Scanner:
    """A class used to represent the Web Scanner

    Attributes
    ----------
    url: str
        A string representing the url of the website to be scanned

    Methods
    -------
    scanner()
        Scans Google News using the url of the class instance.
        Returns a list with all the headlines of the news.
    keyword_search(keyword)
        Scans Google News using a keyword to filter the results.
        Returns a List with headlines and links of the news
        containing the keyword.
    most_common_words()
        Scans Google News and return a List of Tuples with
        a word and an integer representing how many times the
        word has been found in the headlines of the news.
    """

    def __init__(self, url):
        self.url = url

    def scanner(self) -> List[str]:
        titles = []

        page = requests.get(self.url)

        soup = BeautifulSoup(page.content, 'html.parser')
        # results = soup.find(class_='VDXfz')

        news_elems = soup.find_all('h3', class_='ipQwMb ekueJc gEATFF RD0gLb')
        # pprint(page.content)
        for news_elem in news_elems:
            # print(news_elem, end='\n'*2)
            title_elem = news_elem.find('a', class_='DY5T1d')
            # print(title_elem.text)
            titles.append(title_elem.text)

        return titles

    def keyword_search(self, keyword: Any) -> List[str]:

        page = requests.get(self.url)

        soup = BeautifulSoup(page.content, 'html.parser')

        news_links = soup.find_all('h3', string=lambda text: keyword in text.lower())

        links_list = []

        for n_link in news_links:
            link = n_link.find('a')['href']
            links_list.append((n_link.text.strip()))
            links_list.append(f"Read here: https://news.google.com{link}")

        return links_list

    def most_common_words(self) -> List[Tuple[str, int]]:

        page = requests.get(self.url)

        soup = BeautifulSoup(page.content, 'html.parser')

        news_links = soup.find_all('h3')

        headlines_list = []

        stop_words = nltk.corpus.stopwords.words('english')

        for n_link in news_links:
            headlines_list.append(nltk.word_tokenize(n_link.text.strip()))

        filtered_list = [w.lower() for sublist in headlines_list for w in sublist]

        black_list = [".", ",", ":", "\"", "!", "?", "\'", "*", "'s", "'s", "the", '-', 'the', '’', '‘',
                      "n't"] + stop_words

        for word in filtered_list:
            if word in black_list:
                filtered_list.remove(word)

        return Counter(filtered_list).most_common(15)


def results_docx(arr: List[str]) -> None:
    """Gets a List of strings from the keyword_search()
    method and creates a .docx file containing the news
    headlines and links.
    Parameters
    ----------
    arr: List[str]
        The list with the headlines and links
    """

    mydoc = Document()

    for i in range(0, len(arr) - 1, 2):
        mydoc.add_heading(arr[i], 1)
        mydoc.add_paragraph(arr[i + 1])

    mydoc.save('results.docx')


def category_picker(topic: int) -> str:
    """Provided an integer from 1-10 returns
    the link to the category picked by the user.

    Parameters
    ----------
    topic: int
        The number of the topic.

    Returns
    -------
    str
        A string containing the link
    """
    switcher = {
        1: 'https://news.google.com/topstories',
        2: 'https://news.google.com/topics/CAAqIggKIhxDQkFTRHdvSkwyMHZNR1F3TmpCbkVnSmxiaWdBUAE',
        3: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRGx1YlY4U0FtVnVHZ0pEUVNnQVAB',
        4: 'https://news.google.com/topics/CAAqHAgKIhZDQklTQ2pvSWJHOWpZV3hmZGpJb0FBUAE',
        5: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRGx6TVdZU0FtVnVHZ0pWVXlnQVAB',
        6: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRGRqTVhZU0FtVnVHZ0pWVXlnQVAB',
        7: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNREpxYW5RU0FtVnVHZ0pWVXlnQVAB',
        8: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRFp1ZEdvU0FtVnVHZ0pWVXlnQVAB',
        9: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRFp0Y1RjU0FtVnVHZ0pWVXlnQVAB',
        10: 'https://news.google.com/topics/CAAqIQgKIhtDQkFTRGdvSUwyMHZNR3QwTlRFU0FtVnVLQUFQAQ',
    }
    return switcher.get(topic, "Wrong topic!")


def input_validator(prompt: Any) -> int:
    """Validates the input from the user to be
    an int from 1-10.

    Parameters
    ----------
    prompt: str
        Asks the user for an input

    Returns
    -------
    int
        An integer representing the topic, which the user picked
    """

    while True:
        value = input(prompt)
        if value not in ('1', '2', '3', '4', '5', '6', '7', '8', '9', '10'):
            print("Sorry, your input is not from 1 to 10. Pick again:")
            continue
        else:
            break
    return int(value)


def main():
    categories = {
        1: "Top Stories",
        2: "Country(Users own country)",
        3: "World",
        4: "Your Local News",
        5: "Business",
        6: "Technology",
        7: "Entertainment",
        8: "Sports",
        9: "Science",
        10: "Health",
    }

    print("Welcome to the news analyzer".title().strip() + '\n')
    pprint(categories)
    print("\n")
    category = input_validator("Which category would you like to read about (1-10): ")
    print("You picked the topic " + categories[category] + "\n")
    print("Please wait for the program to crawl the web! \n")
    # time.sleep(3)
    my_scanner = Scanner(category_picker(category))
    # pprint(myScanner.scanner())
    print(my_scanner.most_common_words())
    user_keyword = input("A keyword you want links for: ")
    results_docx(my_scanner.keyword_search(user_keyword))
    print("Check your folder for the results.dock file to see the headlines and links.")


if __name__ == "__main__":
    main()
