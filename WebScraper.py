import urllib.request
from docx import Document
from bs4 import BeautifulSoup
import requests
from pprint import pprint
import time
import nltk
from collections import Counter
# Uncomment the lines below if you get an error on the first run of the program. After that you can delete it.
# nltk.download('punkt')
# nltk.download('stopwords')


class Scanner:

    def __init__(self, url):
        self.url = url

    def scanner(self):
        titles = []

        page = requests.get(self.url)

        soup = BeautifulSoup(page.content, 'html.parser')
        # results = soup.find(class_='VDXfz')

        news_elems = soup.find_all('h3', class_='ipQwMb ekueJc gEATFF RD0gLb')
        # pprint(page.content)
        for news_elem in news_elems:
            # print(news_elem, end='\n'*2)
            title_elem = news_elem.find('a', class_='DY5T1d')
            # print(title_elem.text)
            titles.append(title_elem.text)

        return titles

    def keyword_search(self, keyword):

        page = requests.get(self.url)

        soup = BeautifulSoup(page.content, 'html.parser')

        news_links = soup.find_all('h3', string=lambda text: keyword in text.lower())

        links_list = []

        for n_link in news_links:
            link = n_link.find('a')['href']
            links_list.append((n_link.text.strip()))
            links_list.append(f"Read here: https://news.google.com{link}")

        return links_list

    def most_common_words(self):

        page = requests.get(self.url)

        soup = BeautifulSoup(page.content, 'html.parser')

        news_links = soup.find_all('h3')

        headlines_list = []

        stop_words = set(nltk.corpus.stopwords.words('english'))

        for n_link in news_links:
            headlines_list.append(nltk.word_tokenize(n_link.text.strip()))

        filtered_list = [w.lower() for sublist in headlines_list for w in sublist if w not in stop_words]
        return Counter(filtered_list).most_common(25)


def results_docx(arr):

    mydoc = Document()

    for i in range(0, len(arr) - 1, 2):
        mydoc.add_heading(arr[i], 1)
        mydoc.add_paragraph(arr[i + 1])

    mydoc.save('results.docx')


def category_picker(topic):
    switcher = {
        1: 'https://news.google.com/topstories',
        2: 'https://news.google.com/topics/CAAqIggKIhxDQkFTRHdvSkwyMHZNR1F3TmpCbkVnSmxiaWdBUAE',
        3: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRGx1YlY4U0FtVnVHZ0pEUVNnQVAB',
        4: 'https://news.google.com/topics/CAAqHAgKIhZDQklTQ2pvSWJHOWpZV3hmZGpJb0FBUAE',
        5: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRGx6TVdZU0FtVnVHZ0pWVXlnQVAB',
        6: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRGRqTVhZU0FtVnVHZ0pWVXlnQVAB',
        7: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNREpxYW5RU0FtVnVHZ0pWVXlnQVAB',
        8: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRFp1ZEdvU0FtVnVHZ0pWVXlnQVAB',
        9: 'https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRFp0Y1RjU0FtVnVHZ0pWVXlnQVAB',
        10: 'https://news.google.com/topics/CAAqIQgKIhtDQkFTRGdvSUwyMHZNR3QwTlRFU0FtVnVLQUFQAQ',
    }
    return switcher.get(topic, "Wrong topic!")


def input_validator(prompt):
    while True:
        value = input(prompt)
        if value not in ('1', '2', '3', '4', '5', '6', '7', '8', '9', '10'):
            print("Sorry, your input is not from 1 to 8. Pick again:")
            continue
        else:
            break
    return int(value)


categories = {
    1: "Top Stories",
    2: "Country(Users own country)",
    3: "World",
    4: "Your Local News",
    5: "Business",
    6: "Technology",
    7: "Entertainment",
    8: "Sports",
    9: "Science",
    10: "Health",
}

print("Welcome to the news analyzer".title().strip())
pprint(categories)
category = input_validator("Which category would you like to read about (1-8):")
print("You picked the topic " + categories[category])
print("Please wait for the program to crawl the web!")
# time.sleep(3)
myScanner = Scanner(category_picker(category))
# pprint(myScanner.scanner())
#user_keyword = input("A keyword you want links for:")
#results_docx(myScanner.keyword_search(user_keyword))
print(myScanner.most_common_words())
